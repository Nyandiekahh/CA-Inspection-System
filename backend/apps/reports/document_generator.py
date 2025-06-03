# apps/reports/document_generator.py - COMPLETE IMPLEMENTATION
import os
import math
from datetime import datetime
from typing import Dict, List, Any, Optional
from io import BytesIO

# Document generation libraries - REMOVED PDF IMPORTS
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.enum.dml import MSO_THEME_COLOR_INDEX

from django.conf import settings
from django.core.files.base import ContentFile
from django.template import Template, Context

from .models import InspectionReport, ReportImage, ERPCalculation

class ProfessionalDocumentGenerator:
    """Professional DOCX document generator for CA inspection reports"""
    
    def __init__(self, report: InspectionReport):
        self.report = report
        self.inspection = report.inspection
        self.broadcaster = self.inspection.broadcaster
        
        # Image categories mapping to match frontend
        self.image_categories = {
            'site_overview': 'Site Overview',
            'tower_mast': 'Tower/Mast Structure',
            'transmitter_equipment': 'Transmitter Equipment',
            'antenna': 'Antenna System',
            'studio_transmitter_link': 'Studio to Transmitter Link',
            'filter_equipment': 'Filter Equipment',
            'other_equipment': 'Other Equipment'
        }
    
    def generate_documents(self, formats: List[str] = None) -> Dict[str, str]:
        """Generate professional documents - DOCX ONLY"""
        if not formats:
            formats = ['docx']
        
        results = {}
        
        try:
            # REMOVED: PDF generation - only DOCX
            if 'docx' in formats:
                docx_path = self.generate_professional_docx()
                results['docx'] = docx_path
            
            # Update report status
            self.report.status = 'completed'
            self.report.date_completed = datetime.now()
            self.report.save()
            
            return results
            
        except Exception as e:
            print(f"Document generation failed: {str(e)}")
            self.report.status = 'draft'
            self.report.save()
            raise
    
    def generate_professional_docx(self) -> str:
        """Generate professional Word document matching CA templates"""
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = self.report.title
        doc.core_properties.author = self.inspection.inspector.get_full_name()
        doc.core_properties.created = self.report.created_at
        
        # Build document content
        self._build_docx_header(doc)
        self._build_docx_findings_section(doc)
        self._build_docx_conclusions_section(doc)
        self._build_docx_signature(doc)
        
        # Save document
        buffer = BytesIO()
        doc.save(buffer)
        
        filename = f"{self.report.reference_number.replace('/', '_')}.docx"
        self.report.generated_docx.save(
            filename,
            ContentFile(buffer.getvalue()),
            save=True
        )
        
        return self.report.generated_docx.path
    
    def _build_docx_header(self, doc: Document):
        """Build document header matching CA format"""
        # Reference number - bold, left aligned
        ref_para = doc.add_paragraph()
        ref_run = ref_para.add_run(self.report.reference_number)
        ref_run.bold = True
        ref_run.font.size = Pt(12)
        
        # Date
        date_str = self._format_date_with_suffix(self.inspection.inspection_date)
        doc.add_paragraph(date_str)
        doc.add_paragraph()  # Empty line
        
        # Addressee - bold
        to_para = doc.add_paragraph()
        to_run = to_para.add_run("TO: D/MIRC")
        to_run.bold = True
        
        thru_para = doc.add_paragraph()
        thru_run = thru_para.add_run("THRO': PO/NR/MIRC")
        thru_run.bold = True
        
        doc.add_paragraph()  # Empty line
        
        # Subject line - bold and centered
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(f"RE: {self.report.title}")
        title_run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Empty line
        
        # Reference text
        ref_text = "The above subject matter refers.\n\nReference is made to the above subject."
        doc.add_paragraph(ref_text)
        
        # Inspection details
        inspector_name = self.inspection.inspector.get_full_name()
        inspection_date = self._format_date_with_suffix(self.inspection.inspection_date)
        contact_name = self.inspection.contact_name or 'their representative'
        
        inspection_para = doc.add_paragraph()
        inspection_text = f"The transmit station was inspected by MIRC officer {inspector_name} on {inspection_date} in the presence of their representative {contact_name}."
        inspection_para.add_run(inspection_text)
        
        doc.add_paragraph()  # Empty line
        
        # FINDINGS header - bold
        findings_para = doc.add_paragraph()
        findings_run = findings_para.add_run("FINDINGS")
        findings_run.bold = True
        findings_run.font.size = Pt(14)
    
    def _build_docx_findings_section(self, doc: Document):
        """Build findings section with all technical details"""
        
        # A. SITE
        self._add_section_header(doc, "A. SITE")
        self._build_site_table(doc)
        self._add_section_images(doc, 'site_overview')
        doc.add_paragraph()
        
        # B. MAST/TOWER
        self._add_section_header(doc, "B. MAST")
        self._build_tower_table(doc)
        self._add_section_images(doc, 'tower_mast')
        doc.add_paragraph()
        
        # C. TRANSMITTER
        self._add_section_header(doc, "C. TRANSMITTER")
        if self.inspection.station_type == 'TV' or self.inspection.station_type == 'DTT':
            self._build_tv_transmitter_table(doc)
        else:
            self._build_fm_transmitter_table(doc)
        self._add_section_images(doc, 'transmitter_equipment')
        doc.add_paragraph()
        
        # D. ANTENNA SYSTEM
        self._add_section_header(doc, "D. ANTENNA SYSTEM")
        self._build_antenna_table(doc)
        self._add_section_images(doc, 'antenna')
        doc.add_paragraph()
        
        # E. FILTER (if applicable)
        if self.inspection.filter_manufacturer or self._has_images('filter_equipment'):
            self._add_section_header(doc, "E. FILTER")
            self._build_filter_table(doc)
            self._add_section_images(doc, 'filter_equipment')
            doc.add_paragraph()
        
        # F. STUDIO TO TRANSMITTER LINK (if applicable)
        if self.inspection.studio_manufacturer or self._has_images('studio_transmitter_link'):
            self._add_section_header(doc, "F. STUDIO TO TRANSMITTER LINK")
            self._build_stl_table(doc)
            self._add_section_images(doc, 'studio_transmitter_link')
            doc.add_paragraph()
        
        # G. ERP CALCULATION - UPDATED TO FETCH FROM INSPECTION
        self._add_section_header(doc, "G. ERP CALCULATION")
        self._build_erp_section(doc)
        doc.add_paragraph()
    
    def _build_site_table(self, doc: Document):
        """Build site information table"""
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        table.autofit = False
        
        # Set column widths
        table.columns[0].width = Inches(2.5)
        table.columns[1].width = Inches(4.0)
        
        # Site data
        site_data = [
            ("Name:", self.inspection.transmitting_site_name or 'Not specified'),
            ("Coordinates:", f"{self.inspection.longitude or ''}\n{self.inspection.latitude or ''}"),
            ("Elevation:", f"{self.inspection.altitude or 'Not specified'} M")
        ]
        
        for i, (label, value) in enumerate(site_data):
            # Label cell - bold
            label_cell = table.cell(i, 0)
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.bold = True
            
            # Value cell
            value_cell = table.cell(i, 1)
            value_cell.text = value
    
    def _build_tower_table(self, doc: Document):
        """Build tower/mast information table"""
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        # Tower data
        tower_type = self.inspection.get_tower_type_display() if self.inspection.tower_type else 'Not specified'
        height = f"{self.inspection.height_above_ground or 'Not specified'}M"
        
        tower_data = [
            ("Type:", tower_type),
            ("Height:", height)
        ]
        
        for i, (label, value) in enumerate(tower_data):
            # Label cell - bold
            label_cell = table.cell(i, 0)
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.bold = True
            
            # Value cell
            value_cell = table.cell(i, 1)
            value_cell.text = value
    
    def _build_fm_transmitter_table(self, doc: Document):
        """Build FM transmitter table with equipment"""
        
        # Exciter section
        exciter_para = doc.add_paragraph()
        exciter_run = exciter_para.add_run("EXCITER")
        exciter_run.bold = True
        
        exciter_table = doc.add_table(rows=5, cols=2)
        exciter_table.style = 'Table Grid'
        
        exciter_data = [
            ("Make:", self.inspection.exciter_manufacturer or 'Not Seen'),
            ("Model:", self.inspection.exciter_model_number or 'Not Seen'),
            ("Serial:", self.inspection.exciter_serial_number or 'Not Seen'),
            ("Nominal Power:", f"{self.inspection.exciter_nominal_power or 'Not Seen'} W"),
            ("Power Output:", f"{self.inspection.exciter_actual_reading or 'Not Seen'} W")
        ]
        
        for i, (label, value) in enumerate(exciter_data):
            # Label cell - bold
            label_cell = exciter_table.cell(i, 0)
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.bold = True
            
            # Value cell
            value_cell = exciter_table.cell(i, 1)
            value_cell.text = value
        
        doc.add_paragraph()
        
        # Amplifier section
        amp_para = doc.add_paragraph()
        amp_run = amp_para.add_run("AMPLIFIER")
        amp_run.bold = True
        
        amp_table = doc.add_table(rows=5, cols=2)
        amp_table.style = 'Table Grid'
        
        amp_data = [
            ("Make:", self.inspection.amplifier_manufacturer or 'Not Seen'),
            ("Model:", self.inspection.amplifier_model_number or 'Not Seen'),
            ("Serial:", self.inspection.amplifier_serial_number or 'Not Seen'),
            ("Nominal Power:", f"{self.inspection.amplifier_nominal_power or 'Not Seen'} W"),
            ("Power Output:", f"{self.inspection.amplifier_actual_reading or 'Not Seen'} W")
        ]
        
        for i, (label, value) in enumerate(amp_data):
            # Label cell - bold
            label_cell = amp_table.cell(i, 0)
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.bold = True
            
            # Value cell
            value_cell = amp_table.cell(i, 1)
            value_cell.text = value
        
        # Frequency
        doc.add_paragraph()
        freq_para = doc.add_paragraph()
        freq_run = freq_para.add_run(f"Frequency: {self.inspection.transmit_frequency or 'Not Specified'}")
        freq_run.bold = True
    
    def _build_tv_transmitter_table(self, doc: Document):
        """Build TV transmitter table for multiple channels - UPDATED"""
        # Use the improved method for better formatting
        self._build_tv_transmitter_table_improved(doc)
    
    def _build_antenna_table(self, doc: Document):
        """Build antenna system table"""
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        antenna_data = [
            ("Manufacturer:", self.inspection.antenna_manufacturer or 'Not specified'),
            ("Model No.:", self.inspection.antenna_model_number or 'Not specified'),
            ("Type:", self.inspection.antenna_type or 'Not specified'),
            ("Polarization:", self.inspection.get_polarization_display() if self.inspection.polarization else 'Not specified'),
            ("Gain:", f"{self.inspection.antenna_gain or 'Not specified'} dBi"),
            ("Height on the Tower:", f"{self.inspection.height_on_tower or 'Not specified'}M")
        ]
        
        for i, (label, value) in enumerate(antenna_data):
            # Label cell - bold
            label_cell = table.cell(i, 0)
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.bold = True
            
            # Value cell
            value_cell = table.cell(i, 1)
            value_cell.text = value
    
    def _build_filter_table(self, doc: Document):
        """Build filter equipment table"""
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        filter_data = [
            ("Manufacturer:", self.inspection.filter_manufacturer or 'Not specified'),
            ("Model:", self.inspection.filter_model_number or 'Not Seen'),
            ("S/No.:", self.inspection.filter_serial_number or 'Not Seen'),
            ("Frequency:", self.inspection.filter_frequency or 'Not specified')
        ]
        
        for i, (label, value) in enumerate(filter_data):
            # Label cell - bold
            label_cell = table.cell(i, 0)
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.bold = True
            
            # Value cell
            value_cell = table.cell(i, 1)
            value_cell.text = value
    
    def _build_stl_table(self, doc: Document):
        """Build Studio to Transmitter Link table"""
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        stl_data = [
            ("Manufacturer:", self.inspection.studio_manufacturer or 'Not specified'),
            ("Model:", self.inspection.studio_model_number or 'Not Seen'),
            ("S/No.:", self.inspection.studio_serial_number or 'Not Seen'),
            ("Frequency:", self.inspection.studio_frequency or 'Not specified'),
            ("Description of Signal Reception:", self.inspection.signal_description or 'Not specified')
        ]
        
        for i, (label, value) in enumerate(stl_data):
            # Label cell - bold
            label_cell = table.cell(i, 0)
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.bold = True
            
            # Value cell
            value_cell = table.cell(i, 1)
            value_cell.text = value
    
    def _build_erp_section(self, doc: Document):
        """Build ERP calculation section - UPDATED TO HANDLE MULTIPLE CHANNELS"""
        
        # Check if we have multiple channels (TV station)
        erp_calculations = self.report.erp_details.all()
        
        if erp_calculations.count() > 1:
            # Use multi-channel table format like SIGNET report
            self._build_erp_table_multi_channel(doc)
        elif erp_calculations.exists():
            # Single channel with ERP calculation data
            self._build_erp_from_calculations(doc, erp_calculations)
        else:
            # Try to get ERP data from inspection first
            erp_kw = self.inspection.effective_radiated_power
            erp_dbw = self.inspection.effective_radiated_power_dbw
            
            if erp_kw or erp_dbw:
                # Use ERP data from inspection
                self._build_erp_from_inspection(doc, erp_kw, erp_dbw)
            else:
                # Last resort: create calculation from inspection data
                self._build_erp_from_equipment_data(doc)
        
        # Authorized ERP
        auth_para = doc.add_paragraph()
        auth_run = auth_para.add_run("Authorized ERP: 10000 W (10 kW)")
        auth_run.bold = True
    
    def _build_erp_from_inspection(self, doc: Document, erp_kw, erp_dbw):
        """Build ERP section using data from inspection record"""
        
        # Channel header
        channel_para = doc.add_paragraph()
        frequency = self.inspection.transmit_frequency or "Unknown"
        channel_run = channel_para.add_run(f"CH.1 ({frequency} MHz)")
        channel_run.bold = True
        
        # ERP data table
        erp_table = doc.add_table(rows=6, cols=2)
        erp_table.style = 'Table Grid'
        
        # Get source data for calculation display
        forward_power = self.inspection.amplifier_actual_reading or self.inspection.exciter_actual_reading or 'Unknown'
        antenna_gain = self.inspection.antenna_gain or 'Unknown'
        
        # Use estimated losses if available, otherwise default
        system_losses = (
            self.inspection.estimated_system_losses or 
            self._calculate_total_losses() or 
            1.5
        )
        
        erp_data = [
            ("Forward Power:", f"{forward_power} W"),
            ("Antenna Gain:", f"{antenna_gain} dBi"),
            ("System Losses:", f"{system_losses} dB"),
            ("ERP Calculation:", "ERP = 10log P(W) + G (dBi) – L (dB)"),
            ("Result (kW):", f"{erp_kw} kW" if erp_kw else 'Not calculated'),
            ("Result (dBW):", f"{erp_dbw} dBW" if erp_dbw else 'Not calculated')
        ]
        
        for i, (label, value) in enumerate(erp_data):
            # Label cell - bold
            label_cell = erp_table.cell(i, 0)
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.bold = True
            
            # Value cell
            value_cell = erp_table.cell(i, 1)
            value_cell.text = str(value)
        
        doc.add_paragraph()
    
    def _build_erp_from_calculations(self, doc: Document, erp_calculations):
        """Build ERP section using ERPCalculation records"""
        
        for calc in erp_calculations:
            # Channel header
            channel_para = doc.add_paragraph()
            channel_run = channel_para.add_run(f"{calc.channel_number} ({calc.frequency_mhz} MHz)")
            channel_run.bold = True
            
            # ERP calculation table
            erp_table = doc.add_table(rows=5, cols=2)
            erp_table.style = 'Table Grid'
            
            erp_data = [
                ("Forward Power:", f"{calc.forward_power_w} W"),
                ("Antenna Gain:", f"{calc.antenna_gain_dbd} dBd"),
                ("Losses:", f"{calc.losses_db} dB"),
                ("ERP Calculation:", "ERP = 10log P(W) + G (dBd) – L (dB)"),
                ("Result:", f"ERP = 10log {calc.forward_power_w}(W) + {calc.antenna_gain_dbd} dBd - {calc.losses_db} dB = {calc.erp_dbw} dBW ({calc.erp_kw} kW)")
            ]
            
            for i, (label, value) in enumerate(erp_data):
                # Label cell - bold
                label_cell = erp_table.cell(i, 0)
                label_para = label_cell.paragraphs[0]
                label_run = label_para.add_run(label)
                label_run.bold = True
                
                # Value cell
                value_cell = erp_table.cell(i, 1)
                value_cell.text = value
            
            doc.add_paragraph()
    
    def _build_erp_from_equipment_data(self, doc: Document):
        """Build ERP section using equipment data (fallback)"""
        
        # Channel header
        channel_para = doc.add_paragraph()
        frequency = self.inspection.transmit_frequency or "Unknown"
        channel_run = channel_para.add_run(f"CH.1 ({frequency} MHz)")
        channel_run.bold = True
        
        # Get equipment data
        forward_power = self.inspection.amplifier_actual_reading or self.inspection.exciter_actual_reading
        antenna_gain = self.inspection.antenna_gain
        
        if forward_power and antenna_gain:
            try:
                # Calculate ERP
                power_w = float(forward_power)
                gain_dbi = float(antenna_gain)
                losses_db = self._calculate_total_losses() or 1.5
                
                # ERP = 10*log10(P) + G - L
                erp_dbw = 10 * math.log10(power_w) + gain_dbi - losses_db
                erp_kw = (10 ** (erp_dbw / 10)) / 1000
                
                # ERP calculation table
                erp_table = doc.add_table(rows=5, cols=2)
                erp_table.style = 'Table Grid'
                
                erp_data = [
                    ("Forward Power:", f"{power_w} W"),
                    ("Antenna Gain:", f"{gain_dbi} dBi"),
                    ("Losses:", f"{losses_db} dB"),
                    ("ERP Calculation:", "ERP = 10log P(W) + G (dBi) – L (dB)"),
                    ("Result:", f"ERP = 10log {power_w}(W) + {gain_dbi} dBi - {losses_db} dB = {erp_dbw:.2f} dBW ({erp_kw:.3f} kW)")
                ]
                
                for i, (label, value) in enumerate(erp_data):
                    # Label cell - bold
                    label_cell = erp_table.cell(i, 0)
                    label_para = label_cell.paragraphs[0]
                    label_run = label_para.add_run(label)
                    label_run.bold = True
                    
                    # Value cell
                    value_cell = erp_table.cell(i, 1)
                    value_cell.text = str(value)
                
            except (ValueError, TypeError):
                # If calculation fails, show equipment data only
                self._build_equipment_data_only(doc, forward_power, antenna_gain)
        else:
            # No sufficient data for calculation
            no_data_para = doc.add_paragraph()
            no_data_para.add_run("ERP calculation not available - insufficient equipment data")
        
        doc.add_paragraph()
    
    def _build_equipment_data_only(self, doc: Document, forward_power, antenna_gain):
        """Build equipment data table when ERP cannot be calculated"""
        
        # Equipment data table
        equip_table = doc.add_table(rows=3, cols=2)
        equip_table.style = 'Table Grid'
        
        equip_data = [
            ("Forward Power:", f"{forward_power or 'Not specified'} W"),
            ("Antenna Gain:", f"{antenna_gain or 'Not specified'} dBi"),
            ("ERP Calculation:", "Calculation not available")
        ]
        
        for i, (label, value) in enumerate(equip_data):
            # Label cell - bold
            label_cell = equip_table.cell(i, 0)
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.bold = True
            
            # Value cell
            value_cell = equip_table.cell(i, 1)
            value_cell.text = value
    
    def _calculate_total_losses(self):
        """Calculate total losses from individual loss components"""
        try:
            antenna_losses = float(self.inspection.estimated_antenna_losses or 0)
            feeder_losses = float(self.inspection.estimated_feeder_losses or 0)
            multiplexer_losses = float(self.inspection.estimated_multiplexer_losses or 0)
            
            total = antenna_losses + feeder_losses + multiplexer_losses
            return total if total > 0 else None
            
        except (ValueError, TypeError, AttributeError):
            return None
    
    def _build_docx_conclusions_section(self, doc: Document):
        """Build conclusions section"""
        
        # H. OBSERVATION (if any)
        if self.report.observations or self.inspection.other_observations:
            self._add_section_header(doc, "H. OBSERVATION")
            observations = self.report.observations or self.inspection.other_observations or ''
            if observations:
                # Split observations into bullet points
                obs_lines = [line.strip() for line in observations.split('\n') if line.strip()]
                for obs in obs_lines:
                    if obs:
                        obs_para = doc.add_paragraph(style='List Bullet')
                        obs_para.add_run(obs)
            doc.add_paragraph()
        
        # I. CONCLUSION
        self._add_section_header(doc, "I. CONCLUSION")
        conclusions = self.report.conclusions or self._generate_auto_conclusions()
        if conclusions:
            # Split conclusions into bullet points
            conc_lines = [line.strip() for line in conclusions.split('\n') if line.strip()]
            for conc in conc_lines:
                if conc:
                    conc_para = doc.add_paragraph(style='List Bullet')
                    conc_para.add_run(conc)
        doc.add_paragraph()
        
        # J. RECOMMENDATION
        self._add_section_header(doc, "J. RECOMMENDATION")
        recommendations = self.report.recommendations or self._generate_auto_recommendations()
        if recommendations:
            # Split recommendations into bullet points
            rec_lines = [line.strip() for line in recommendations.split('\n') if line.strip()]
            for rec in rec_lines:
                if rec:
                    rec_para = doc.add_paragraph(style='List Bullet')
                    rec_para.add_run(rec)
    
    def _build_docx_signature(self, doc: Document):
        """Build signature section"""
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Inspector name - bold
        inspector_name = self.inspection.inspector.get_full_name()
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(inspector_name)
        name_run.bold = True
        
        # Title - bold
        title_para = doc.add_paragraph()
        title_run = title_para.add_run("AO/MIRC/NR")
        title_run.bold = True
    
    def _add_section_header(self, doc: Document, header_text: str):
        """Add a section header"""
        header_para = doc.add_paragraph()
        header_run = header_para.add_run(header_text)
        header_run.bold = True
        header_run.font.size = Pt(12)
    
    # ========== MISSING METHODS - NOW IMPLEMENTED ==========
    
    def _add_section_images(self, doc: Document, image_type):
        """Add images for specific sections with proper descriptions"""
        images = self.report.images.filter(image_type=image_type)
        
        if images.exists():
            for image in images:
                try:
                    # Add image to document
                    if image.image and hasattr(image.image, 'path'):
                        # Add some spacing before image
                        doc.add_paragraph()
                        
                        # Add image paragraph
                        img_para = doc.add_paragraph()
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Calculate image width (convert percentage to inches)
                        # Assuming 6.5 inches as max page width (letter size with margins)
                        max_width = 6.5
                        img_width = max_width * (image.width_percentage / 100)
                        
                        # Add the image
                        run = img_para.add_run()
                        run.add_picture(image.image.path, width=Inches(img_width))
                        
                        # Add caption if available
                        if image.caption:
                            caption_para = doc.add_paragraph()
                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption_run = caption_para.add_run(image.caption)
                            caption_run.italic = True
                            caption_run.font.size = Pt(10)
                        
                        # Add spacing after image
                        doc.add_paragraph()
                        
                except Exception as e:
                    print(f"Error adding image {image.id}: {str(e)}")
                    # Add placeholder text instead
                    placeholder_para = doc.add_paragraph()
                    placeholder_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    placeholder_run = placeholder_para.add_run(f"[Image: {image.caption or image.get_image_type_display()}]")
                    placeholder_run.italic = True

    def _has_images(self, image_type):
        """Check if report has images of specified type"""
        return self.report.images.filter(image_type=image_type).exists()

    def _format_date_with_suffix(self, date_obj):
        """Format date with ordinal suffix (e.g., 28th October 2024)"""
        if not date_obj:
            return "Date not specified"
        
        day = date_obj.day
        month = date_obj.strftime('%B')
        year = date_obj.year
        
        # Add ordinal suffix to day
        if 4 <= day <= 20 or 24 <= day <= 30:
            suffix = 'th'
        else:
            suffix = ['st', 'nd', 'rd'][day % 10 - 1]
        
        return f"{day}{suffix} {month} {year}"

    def _get_tv_channels(self):
        """Get TV channel data from inspection - handles multiple channels"""
        channels = []
        
        # Check if we have ERP calculations with multiple channels
        erp_calculations = self.report.erp_details.all()
        if erp_calculations.exists():
            for calc in erp_calculations:
                channels.append({
                    'channel': calc.channel_number,
                    'frequency': calc.frequency_mhz,
                    'power': calc.forward_power_w,
                    'gain': calc.antenna_gain_dbd,
                    'manufacturer': 'See equipment details',
                    'model': 'See equipment details',
                    'serial': 'See equipment details'
                })
        else:
            # Fallback: create single channel from inspection data
            if self.inspection.transmit_frequency:
                channels.append({
                    'channel': 'CH.1',
                    'frequency': self.inspection.transmit_frequency,
                    'power': self.inspection.amplifier_actual_reading or self.inspection.exciter_actual_reading,
                    'gain': self.inspection.antenna_gain or '11.0',
                    'manufacturer': self.inspection.amplifier_manufacturer or self.inspection.exciter_manufacturer,
                    'model': self.inspection.amplifier_model_number or self.inspection.exciter_model_number,
                    'serial': self.inspection.amplifier_serial_number or self.inspection.exciter_serial_number
                })
        
        return channels

    def _generate_auto_conclusions(self):
        """Generate automatic conclusions based on inspection findings"""
        conclusions = []
        
        # Check ERP compliance
        erp_calculations = self.report.erp_details.all()
        authorized_limit = 10.0  # 10 kW standard limit
        
        for calc in erp_calculations:
            if not calc.is_compliant:
                conclusions.append(
                    f"The licensee is operating above the maximum authorized ERP limit of "
                    f"40 dBW (10 kW) by transmitting at {calc.erp_kw} kW ({calc.channel_number})."
                )
        
        # Check for non-type approved equipment
        non_approved_equipment = []
        
        # Check exciter
        if self.inspection.exciter_manufacturer:
            # Add known non-approved manufacturers (this should be configurable)
            non_approved_brands = ['MAXIVA GATEAIR', 'NEC HPB']  # Example list
            if any(brand.lower() in self.inspection.exciter_manufacturer.lower() for brand in non_approved_brands):
                non_approved_equipment.append(f"exciter {self.inspection.exciter_manufacturer} {self.inspection.exciter_model_number or ''}")
        
        # Check amplifier
        if self.inspection.amplifier_manufacturer:
            non_approved_brands = ['MAXIVA GATEAIR']  # Example list
            if any(brand.lower() in self.inspection.amplifier_manufacturer.lower() for brand in non_approved_brands):
                non_approved_equipment.append(f"amplifier {self.inspection.amplifier_manufacturer} {self.inspection.amplifier_model_number or ''}")
        
        if non_approved_equipment:
            conclusions.append(f"The licensee is operating non-type approved transmitter(s): {', '.join(non_approved_equipment)}.")
        
        # Default conclusion if no issues found
        if not conclusions:
            conclusions.append("The station is operating within authorized parameters.")
        
        return '\n'.join(f"• {conclusion}" for conclusion in conclusions)

    def _generate_auto_recommendations(self):
        """Generate automatic recommendations based on findings"""
        recommendations = []
        
        # Check for violations and generate appropriate recommendations
        erp_calculations = self.report.erp_details.all()
        
        # ERP violations
        erp_violations = [calc for calc in erp_calculations if not calc.is_compliant]
        
        if erp_violations:
            # Check if this is a repeat violation (this would need violation history tracking)
            # For now, assume first-time violation
            violation_channels = [calc.channel_number for calc in erp_violations]
            
            if len(violation_channels) > 1:
                recommendations.append(
                    f"The licensee to be issued with notice of violation for exceeding "
                    f"authorized ERP limit of 10kW for {', '.join(violation_channels)}."
                )
            else:
                recommendations.append(
                    f"The licensee to be issued with notice of violation for exceeding "
                    f"authorized ERP limit of 10kW for {violation_channels[0]}."
                )
        
        # Non-type approved equipment
        if self.inspection.exciter_manufacturer:
            non_approved_brands = ['MAXIVA GATEAIR', 'NEC HPB']
            if any(brand.lower() in self.inspection.exciter_manufacturer.lower() for brand in non_approved_brands):
                recommendations.append(
                    "The licensee to be issued with notice of violation for operating "
                    "non-type approved transmitter equipment."
                )
        
        # Equipment maintenance recommendations
        if self.inspection.other_observations:
            if 'rust' in self.inspection.other_observations.lower():
                recommendations.append("The licensee should address tower rust protection issues.")
            
            if 'filter' in self.inspection.other_observations.lower():
                recommendations.append("The licensee should ensure proper filter installation and maintenance.")
        
        # Default recommendation
        if not recommendations:
            recommendations.append("The licensee should continue to maintain the station within authorized parameters.")
        
        return '\n'.join(f"• {recommendation}" for recommendation in recommendations)

    def _build_tv_transmitter_table_improved(self, doc: Document):
        """Build improved TV transmitter table for multiple channels"""
        channels = self._get_tv_channels()
        
        if not channels:
            # Fallback to FM table format
            self._build_fm_transmitter_table(doc)
            return
        
        # Create main table with proper structure
        num_channels = len(channels)
        table = doc.add_table(rows=8, cols=num_channels + 1)  # +1 for labels column
        table.style = 'Table Grid'
        
        # Headers row
        label_cell = table.cell(0, 0)
        label_cell.text = "Channel Freq. (MHz)"
        label_para = label_cell.paragraphs[0]
        label_run = label_para.runs[0]
        label_run.bold = True
        
        # Channel headers
        for i, channel in enumerate(channels, 1):
            header_cell = table.cell(0, i)
            header_cell.text = f"{channel['channel']}\n({channel['frequency']} MHz)"
            header_para = header_cell.paragraphs[0]
            header_run = header_para.runs[0]
            header_run.bold = True
        
        # Equipment rows
        equipment_rows = [
            ("Make:", [ch.get('manufacturer', 'Not Seen') for ch in channels]),
            ("Model:", [ch.get('model', 'Not Seen') for ch in channels]),
            ("S/No.:", [ch.get('serial', 'Not Seen') for ch in channels]),
            ("Nominal Power (W):", [f"{ch.get('power', 'Unknown')} W" for ch in channels]),
            ("Power Output (W):", [f"{ch.get('power', 'Unknown')} W" for ch in channels]),
            ("Gain (dBd):", [f"{ch.get('gain', '11.0')} dBd" for ch in channels]),
        ]
        
        for row_idx, (label, values) in enumerate(equipment_rows, 1):
            # Label cell
            label_cell = table.cell(row_idx, 0)
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.bold = True
            
            # Value cells
            for col_idx, value in enumerate(values, 1):
                if col_idx < len(table.columns):
                    table.cell(row_idx, col_idx).text = str(value)

    def _build_erp_table_multi_channel(self, doc: Document):
        """Build ERP calculation table for multiple channels like SIGNET report"""
        erp_calculations = self.report.erp_details.all()
        
        if not erp_calculations.exists():
            self._build_erp_from_equipment_data(doc)
            return
        
        # Create table with channels as columns
        num_channels = erp_calculations.count()
        table = doc.add_table(rows=6, cols=num_channels + 1)
        table.style = 'Table Grid'
        
        # Headers
        table.cell(0, 0).text = "CHANNEL"
        for i, calc in enumerate(erp_calculations, 1):
            header_cell = table.cell(0, i)
            header_run = header_cell.paragraphs[0].add_run(calc.channel_number)
            header_run.bold = True
        
        # Rows
        row_labels = ["Forward Power:", "Antenna Gain:", "Losses:", "ERP Calculation", "Result"]
        
        for row_idx, label in enumerate(row_labels, 1):
            # Label cell
            label_cell = table.cell(row_idx, 0)
            label_run = label_cell.paragraphs[0].add_run(label)
            label_run.bold = True
            
            # Data cells
            for col_idx, calc in enumerate(erp_calculations, 1):
                cell = table.cell(row_idx, col_idx)
                
                if row_idx == 1:  # Forward Power
                    cell.text = f"{calc.forward_power_w} W"
                elif row_idx == 2:  # Antenna Gain
                    cell.text = f"{calc.antenna_gain_dbd} dBd"
                elif row_idx == 3:  # Losses
                    cell.text = f"{calc.losses_db} dB"
                elif row_idx == 4:  # ERP Calculation
                    cell.text = "ERP=10log P(W) + G (dBd) – L (dB)"
                elif row_idx == 5:  # Result
                    result_text = (
                        f"ERP=10log {calc.forward_power_w}(W) + "
                        f"{calc.antenna_gain_dbd} dBd - {calc.losses_db} dB = "
                        f"{calc.erp_dbw} dBW ({calc.erp_kw} kW)"
                    )
                    cell.text = result_text
        
        doc.add_paragraph()