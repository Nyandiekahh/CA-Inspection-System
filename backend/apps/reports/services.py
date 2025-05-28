# apps/reports/services.py
import os
import math
from datetime import datetime
from typing import Dict, List, Any, Optional
from io import BytesIO
from PIL import Image

# Document generation libraries
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

from django.conf import settings
from django.core.files.base import ContentFile
from django.template import Template, Context

from .models import InspectionReport, ReportImage, ERPCalculation

class DocumentGenerationService:
    """Main service for generating inspection reports"""
    
    def __init__(self, report: InspectionReport):
        self.report = report
        self.inspection = report.inspection
        self.broadcaster = self.inspection.broadcaster
        
    def generate_documents(self, formats: List[str] = None) -> Dict[str, str]:
        """Generate report documents in specified formats"""
        if not formats:
            formats = ['pdf'] if self.report.preferred_format == 'pdf' else ['docx']
        
        results = {}
        
        try:
            if 'pdf' in formats:
                pdf_path = self.generate_pdf()
                results['pdf'] = pdf_path
                
            if 'docx' in formats:
                docx_path = self.generate_docx()
                results['docx'] = docx_path
                
            # Update report status
            self.report.status = 'completed'
            self.report.date_completed = datetime.now()
            self.report.save()
            
            return results
            
        except Exception as e:
            # Log error and update status
            print(f"Document generation failed: {str(e)}")
            self.report.status = 'draft'
            self.report.save()
            raise
    
    def generate_pdf(self) -> str:
        """Generate PDF report"""
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Build story (content)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=14,
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=12,
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        
        # Header
        story.extend(self._build_pdf_header(styles, title_style))
        
        # Site Information
        story.extend(self._build_pdf_site_info(styles, heading_style))
        
        # Mast/Tower Information
        story.extend(self._build_pdf_tower_info(styles, heading_style))
        
        # Transmitter Information
        story.extend(self._build_pdf_transmitter_info(styles, heading_style))
        
        # Antenna System
        story.extend(self._build_pdf_antenna_info(styles, heading_style))
        
        # ERP Calculations
        story.extend(self._build_pdf_erp_calculations(styles, heading_style))
        
        # Observations and Conclusions
        story.extend(self._build_pdf_conclusions(styles, heading_style))
        
        # Build PDF
        doc.build(story)
        
        # Save to file
        pdf_content = buffer.getvalue()
        buffer.close()
        
        filename = f"{self.report.reference_number.replace('/', '_')}.pdf"
        self.report.generated_pdf.save(
            filename,
            ContentFile(pdf_content),
            save=True
        )
        
        return self.report.generated_pdf.path
    
    def generate_docx(self) -> str:
        """Generate Word document"""
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = self.report.title
        doc.core_properties.author = self.inspection.inspector.get_full_name()
        doc.core_properties.created = self.report.created_at
        
        # Build document content
        self._build_docx_header(doc)
        self._build_docx_site_info(doc)
        self._build_docx_tower_info(doc)
        self._build_docx_transmitter_info(doc)
        self._build_docx_antenna_info(doc)
        self._build_docx_erp_calculations(doc)
        self._build_docx_conclusions(doc)
        
        # Save document
        buffer = BytesIO()
        doc.save(buffer)
        
        filename = f"{self.report.reference_number.replace('/', '_')}.docx"
        self.report.generated_docx.save(
            filename,
            ContentFile(buffer.getvalue()),
            save=True
        )
        
        return self.report.generated_docx.path
    
    def _build_pdf_header(self, styles, title_style) -> List:
        """Build PDF header section"""
        story = []
        
        # Reference number
        story.append(Paragraph(f"<b>{self.report.reference_number}</b>", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Date
        date_str = self.inspection.inspection_date.strftime("%d%s %B %Y")
        day = self.inspection.inspection_date.day
        suffix = 'th' if 11 <= day <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
        date_str = self.inspection.inspection_date.strftime(f"%d{suffix} %B %Y")
        story.append(Paragraph(date_str, styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Addressee
        story.append(Paragraph("<b>TO: D/MIRC</b>", styles['Normal']))
        story.append(Paragraph("<b>THRO': PO/NR/MIRC</b>", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Title
        story.append(Paragraph(f"<b>RE: {self.report.title}</b>", title_style))
        story.append(Spacer(1, 20))
        
        # Opening paragraph
        inspector_name = self.inspection.inspector.get_full_name()
        inspection_date = self.inspection.inspection_date.strftime("%-d%s %B %Y")
        day = self.inspection.inspection_date.day
        suffix = 'th' if 11 <= day <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
        inspection_date = self.inspection.inspection_date.strftime(f"%-d{suffix} %B %Y")
        
        opening = f"""The above subject matter refers.
        
Reference is made to the above subject.

The transmit station was inspected by MIRC officer {inspector_name} on {inspection_date} in the presence of their representative {self.inspection.contact_name or 'their representative'}."""
        
        story.append(Paragraph(opening, styles['Normal']))
        story.append(Spacer(1, 20))
        
        # FINDINGS header
        story.append(Paragraph("<b>FINDINGS</b>", styles['Heading1']))
        story.append(Spacer(1, 12))
        
        return story
    
    def _build_pdf_site_info(self, styles, heading_style) -> List:
        """Build site information section"""
        story = []
        
        story.append(Paragraph("<b>A. SITE</b>", heading_style))
        
        # Create site info table
        site_data = [
            ['Name:', self.inspection.transmitting_site_name or 'Not specified'],
            ['Coordinates:', f"{self.inspection.longitude or ''}\n{self.inspection.latitude or ''}"],
            ['Elevation:', f"{self.inspection.altitude or 'Not specified'} M"]
        ]
        
        site_table = Table(site_data, colWidths=[2*inch, 4*inch])
        site_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        
        story.append(site_table)
        story.append(Spacer(1, 20))
        
        return story
    
    def _build_pdf_tower_info(self, styles, heading_style) -> List:
        """Build tower information section"""
        story = []
        
        story.append(Paragraph("<b>B. MAST</b>", heading_style))
        
        # Tower type and height
        tower_type = self.inspection.get_tower_type_display() if self.inspection.tower_type else 'Not specified'
        height = f"{self.inspection.height_above_ground or 'Not specified'}M"
        
        tower_data = [
            ['Type:', tower_type],
            ['Height:', height]
        ]
        
        tower_table = Table(tower_data, colWidths=[2*inch, 4*inch])
        tower_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        
        story.append(tower_table)
        
        # Add tower image if available
        tower_images = self.report.images.filter(image_type='tower_structure')
        if tower_images.exists():
            story.append(Spacer(1, 12))
            for img in tower_images[:1]:  # First tower image
                story.append(self._create_pdf_image(img))
        
        story.append(Spacer(1, 20))
        
        return story
    
    def _build_pdf_transmitter_info(self, styles, heading_style) -> List:
        """Build transmitter information section"""
        story = []
        
        story.append(Paragraph("<b>C. TRANSMITTER</b>", heading_style))
        
        # Build transmitter table with multiple channels/frequencies
        transmitter_data = self._get_transmitter_table_data()
        
        if transmitter_data:
            transmitter_table = Table(transmitter_data)
            transmitter_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            
            story.append(transmitter_table)
        
        # Add equipment images
        equipment_images = self.report.images.filter(
            image_type__in=['exciter', 'amplifier', 'filter']
        )
        if equipment_images.exists():
            story.append(Spacer(1, 12))
            for img in equipment_images:
                story.append(self._create_pdf_image(img))
                story.append(Spacer(1, 8))
        
        story.append(Spacer(1, 20))
        
        return story
    
    def _build_pdf_antenna_info(self, styles, heading_style) -> List:
        """Build antenna system section"""
        story = []
        
        story.append(Paragraph("<b>D. ANTENNA SYSTEM</b>", heading_style))
        
        # Antenna details table
        antenna_data = [
            ['Manufacturer:', self.inspection.antenna_manufacturer or 'Not specified'],
            ['Model No.:', self.inspection.antenna_model_number or 'Not specified'],
            ['Type:', self.inspection.antenna_type or 'Not specified'],
            ['Polarization:', self.inspection.get_polarization_display() if self.inspection.polarization else 'Not specified'],
            ['Gain:', f"{self.inspection.antenna_gain or 'Not specified'} dBd"],
            ['Height on the Tower:', f"{self.inspection.height_on_tower or 'Not specified'}M"]
        ]
        
        antenna_table = Table(antenna_data, colWidths=[2.5*inch, 3.5*inch])
        antenna_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        
        story.append(antenna_table)
        
        # Add antenna images
        antenna_images = self.report.images.filter(image_type='antenna_system')
        if antenna_images.exists():
            story.append(Spacer(1, 12))
            for img in antenna_images:
                story.append(self._create_pdf_image(img))
        
        story.append(Spacer(1, 20))
        
        return story
    
    def _build_pdf_erp_calculations(self, styles, heading_style) -> List:
        """Build ERP calculations section"""
        story = []
        
        story.append(Paragraph("<b>E. ERP CALCULATION</b>", heading_style))
        
        # Get or create ERP calculations
        erp_calculations = self.report.erp_details.all()
        if not erp_calculations.exists():
            # Create default calculation
            erp_calc = self._create_default_erp_calculation()
            if erp_calc:
                erp_calculations = [erp_calc]
        
        if erp_calculations:
            for calc in erp_calculations:
                # Individual calculation
                calc_text = f"""
<b>{calc.channel_number} ({calc.frequency_mhz} MHz)</b><br/>
<b>Forward Power:</b> {calc.forward_power_w} W<br/>
<b>Antenna Gain:</b> {calc.antenna_gain_dbd} dBd<br/>
<b>Losses:</b> {calc.losses_db} dB<br/>
<br/>
<b>ERP Calculation</b><br/>
<b>ERP=10log P(W) + G (dBd) – L (dB)</b><br/>
ERP=10log {calc.forward_power_w}(W) + {calc.antenna_gain_dbd} dBd - {calc.losses_db} dB = {calc.erp_dbw} dBW ({calc.erp_kw} kW)
"""
                
                story.append(Paragraph(calc_text, styles['Normal']))
                story.append(Spacer(1, 12))
        
        # Authorized ERP
        auth_erp = "10000 W (10kW)"  # Default
        story.append(Paragraph(f"<b>Authorized ERP:</b> {auth_erp}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        return story
    
    def _build_pdf_conclusions(self, styles, heading_style) -> List:
        """Build conclusions and recommendations"""
        story = []
        
        # Observations
        if self.report.observations:
            story.append(Paragraph("<b>H. OBSERVATION</b>", heading_style))
            story.append(Paragraph(self.report.observations, styles['Normal']))
            story.append(Spacer(1, 12))
        
        # Conclusions
        story.append(Paragraph("<b>I. CONCLUSION</b>", heading_style))
        conclusions = self.report.conclusions or self._generate_auto_conclusions()
        story.append(Paragraph(conclusions, styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Recommendations
        story.append(Paragraph("<b>J. RECOMMENDATION</b>", heading_style))
        recommendations = self.report.recommendations or self._generate_auto_recommendations()
        story.append(Paragraph(recommendations, styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Signature
        inspector_name = self.inspection.inspector.get_full_name()
        story.append(Paragraph(f"<b>{inspector_name}</b>", styles['Normal']))
        story.append(Paragraph("<b>AO/MIRC/NR</b>", styles['Normal']))
        
        return story
    
    def _create_pdf_image(self, report_image: ReportImage) -> RLImage:
        """Create PDF image element"""
        try:
            # Open and resize image
            img_path = report_image.image.path
            pil_img = Image.open(img_path)
            
            # Calculate size (maintain aspect ratio)
            max_width = 4 * inch
            max_height = 3 * inch
            
            img_width, img_height = pil_img.size
            ratio = min(max_width/img_width, max_height/img_height)
            
            new_width = img_width * ratio
            new_height = img_height * ratio
            
            # Create ReportLab image
            rl_img = RLImage(img_path, width=new_width, height=new_height)
            
            return rl_img
            
        except Exception as e:
            print(f"Error creating PDF image: {e}")
            return None
    
    # Word Document Generation Methods
    def _build_docx_header(self, doc: Document):
        """Build Word document header"""
        # Reference number
        ref_para = doc.add_paragraph()
        ref_run = ref_para.add_run(self.report.reference_number)
        ref_run.bold = True
        
        # Date
        date_str = self.inspection.inspection_date.strftime("%-d")
        day = self.inspection.inspection_date.day
        suffix = 'th' if 11 <= day <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
        date_str = self.inspection.inspection_date.strftime(f"%-d{suffix} %B %Y")
        
        doc.add_paragraph(date_str)
        doc.add_paragraph()
        
        # Addressee
        to_para = doc.add_paragraph()
        to_run = to_para.add_run("TO: D/MIRC")
        to_run.bold = True
        
        thru_para = doc.add_paragraph()
        thru_run = thru_para.add_run("THRO': PO/NR/MIRC")
        thru_run.bold = True
        
        doc.add_paragraph()
        
        # Title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(f"RE: {self.report.title}")
        title_run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Opening text
        opening_text = f"""The above subject matter refers.

Reference is made to the above subject.

The transmit station was inspected by MIRC officer {self.inspection.inspector.get_full_name()} on {date_str} in the presence of their representative {self.inspection.contact_name or 'their representative'}."""
        
        doc.add_paragraph(opening_text)
        doc.add_paragraph()
        
        # FINDINGS header
        findings_para = doc.add_paragraph()
        findings_run = findings_para.add_run("FINDINGS")
        findings_run.bold = True
        findings_run.font.size = Pt(14)
    
    def _build_docx_site_info(self, doc: Document):
        """Build site information section for Word doc"""
        # Site header
        site_para = doc.add_paragraph()
        site_run = site_para.add_run("A. SITE")
        site_run.bold = True
        
        # Site table
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        # Table data
        table.cell(0, 0).text = "Name:"
        table.cell(0, 1).text = self.inspection.transmitting_site_name or 'Not specified'
        
        table.cell(1, 0).text = "Coordinates:"
        coord_text = f"{self.inspection.longitude or ''}\n{self.inspection.latitude or ''}"
        table.cell(1, 1).text = coord_text
        
        table.cell(2, 0).text = "Elevation:"
        table.cell(2, 1).text = f"{self.inspection.altitude or 'Not specified'} M"
        
        # Make first column bold
        for row in table.rows:
            row.cells[0].paragraphs[0].runs[0].bold = True
    
    def _build_docx_transmitter_info(self, doc: Document):
        """Build transmitter section for Word doc"""
        # Transmitter header
        trans_para = doc.add_paragraph()
        trans_run = trans_para.add_run("C. TRANSMITTER")
        trans_run.bold = True
        
        # Build transmitter table
        transmitter_data = self._get_transmitter_table_data()
        if transmitter_data and len(transmitter_data) > 1:
            rows = len(transmitter_data)
            cols = len(transmitter_data[0])
            
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            
            for i, row_data in enumerate(transmitter_data):
                for j, cell_data in enumerate(row_data):
                    table.cell(i, j).text = str(cell_data)
                    if i == 0:  # Header row
                        table.cell(i, j).paragraphs[0].runs[0].bold = True
        
        # Add equipment images
        equipment_images = self.report.images.filter(
            image_type__in=['exciter', 'amplifier', 'filter']
        )
        for img in equipment_images:
            self._add_docx_image(doc, img)
    
    def _add_docx_image(self, doc: Document, report_image: ReportImage):
        """Add image to Word document"""
        try:
            doc.add_paragraph()
            paragraph = doc.add_paragraph()
            
            # Calculate image size
            max_width = Inches(5)
            run = paragraph.add_run()
            run.add_picture(report_image.image.path, width=max_width)
            
            # Add caption
            if report_image.caption:
                caption_para = doc.add_paragraph(report_image.caption)
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_para.runs[0]
                caption_run.italic = True
                
        except Exception as e:
            print(f"Error adding image to Word doc: {e}")
    
    # Helper Methods
    def _get_transmitter_table_data(self) -> List[List[str]]:
        """Get transmitter data for table generation"""
        # This would typically handle multiple channels/frequencies
        # For now, create a basic table structure
        
        headers = ["Channel Freq. (MHz)", "Make", "Model", "S/No.", "Power (W)", "Gain (dBd)"]
        
        # Basic row with inspection data
        freq = self.inspection.transmit_frequency or "Not specified"
        exciter_make = self.inspection.exciter_manufacturer or "Not specified"
        exciter_model = self.inspection.exciter_model_number or "Not specified"
        exciter_serial = self.inspection.exciter_serial_number or "Not seen"
        power = self.inspection.amplifier_actual_reading or "Not specified"
        gain = self.inspection.antenna_gain or "11.0"
        
        data_row = [freq, exciter_make, exciter_model, exciter_serial, power, gain]
        
        return [headers, data_row]
    
    def _create_default_erp_calculation(self) -> Optional[ERPCalculation]:
        """Create default ERP calculation from inspection data"""
        try:
            # Extract power and frequency info
            forward_power = float(self.inspection.amplifier_actual_reading or 3000)
            antenna_gain = float(self.inspection.antenna_gain or 11.0)
            frequency = self.inspection.transmit_frequency or "Unknown"
            
            # Create calculation
            calc = ERPCalculation.objects.create(
                report=self.report,
                channel_number="CH.1",
                frequency_mhz=frequency,
                forward_power_w=forward_power,
                antenna_gain_dbd=antenna_gain,
                losses_db=1.5
            )
            
            return calc
            
        except (ValueError, TypeError):
            return None
    
    def _generate_auto_conclusions(self) -> str:
        """Generate automatic conclusions based on violations"""
        violations = self.report.violations_found
        
        if not violations:
            return "The station is operating in compliance with the authorized parameters."
        
        # Check for ERP violations
        erp_violations = [v for v in violations if 'ERP' in v.get('type', '')]
        type_approval_violations = [v for v in violations if 'type approval' in v.get('type', '')]
        
        conclusions = []
        
        if erp_violations:
            conclusions.append("The licensee is operating above the maximum authorized ERP limit.")
        
        if type_approval_violations:
            conclusions.append("The licensee is operating non-type approved transmitters.")
        
        return " ".join(conclusions)
    
    def _generate_auto_recommendations(self) -> str:
        """Generate automatic recommendations based on violations"""
        violations = self.report.violations_found
        
        if not violations:
            return "Continue operating within authorized parameters and maintain equipment in good condition."
        
        recommendations = []
        
        # Check for ERP violations
        erp_violations = [v for v in violations if 'ERP' in v.get('type', '')]
        if erp_violations:
            recommendations.append("The licensee to be issued with a notice of violation for exceeding authorized ERP limit.")
        
        # Check for type approval violations
        type_approval_violations = [v for v in violations if 'type approval' in v.get('type', '')]
        if type_approval_violations:
            recommendations.append("The licensee to be issued with a notice of violation for operating non-type approved transmitters.")
        
        return " ".join(recommendations)

class ERPCalculationService:
    """Service for ERP calculations"""
    
    @staticmethod
    def calculate_erp(forward_power_w: float, antenna_gain_dbd: float, losses_db: float = 1.5) -> Dict[str, float]:
        """Calculate ERP values"""
        # ERP = 10*log10(P) + G - L
        erp_dbw = 10 * math.log10(forward_power_w) + antenna_gain_dbd - losses_db
        erp_kw = (10 ** (erp_dbw / 10)) / 1000
        
        return {
            'erp_dbw': round(erp_dbw, 2),
            'erp_kw': round(erp_kw, 3)
        }
    
    @staticmethod
    def check_compliance(erp_kw: float, authorized_kw: float = 10.0) -> Dict[str, Any]:
        """Check ERP compliance"""
        is_compliant = erp_kw <= authorized_kw
        excess_kw = max(0, erp_kw - authorized_kw) if not is_compliant else 0
        
        return {
            'is_compliant': is_compliant,
            'excess_kw': round(excess_kw, 3),
            'percentage_over': round((excess_kw / authorized_kw) * 100, 1) if excess_kw > 0 else 0
        }

class ViolationDetectionService:
    """Service for detecting violations in inspection data"""
    
    def __init__(self, inspection):
        self.inspection = inspection
    
    def detect_violations(self) -> List[Dict[str, Any]]:
        """Detect all violations in the inspection"""
        violations = []
        
        # Check ERP violations
        violations.extend(self._check_erp_violations())
        
        # Check type approval violations
        violations.extend(self._check_type_approval_violations())
        
        # Check safety violations
        violations.extend(self._check_safety_violations())
        
        return violations
    
    def _check_erp_violations(self) -> List[Dict[str, Any]]:
        """Check for ERP limit violations"""
        violations = []
        
        try:
            forward_power = float(self.inspection.amplifier_actual_reading or 0)
            antenna_gain = float(self.inspection.antenna_gain or 11.0)
            
            if forward_power > 0:
                erp_calc = ERPCalculationService.calculate_erp(forward_power, antenna_gain)
                compliance = ERPCalculationService.check_compliance(erp_calc['erp_kw'])
                
                if not compliance['is_compliant']:
                    violations.append({
                        'type': 'ERP_VIOLATION',
                        'severity': 'major',
                        'description': f"Operating above authorized ERP limit by {compliance['excess_kw']} kW",
                        'measured_value': erp_calc['erp_kw'],
                        'authorized_value': 10.0,
                        'excess': compliance['excess_kw']
                    })
        
        except (ValueError, TypeError):
            pass
        
        return violations
    
    def _check_type_approval_violations(self) -> List[Dict[str, Any]]:
        """Check for type approval violations"""
        violations = []
        
        # List of known non-approved equipment (this would be from a database)
        non_approved_equipment = [
            'MAXIVA GATEAIR XTE',
            'NEC HPB-1210',
            # Add more as needed
        ]
        
        # Check exciter
        exciter_model = f"{self.inspection.exciter_manufacturer or ''} {self.inspection.exciter_model_number or ''}".strip()
        if exciter_model in non_approved_equipment:
            violations.append({
                'type': 'TYPE_APPROVAL_VIOLATION',
                'severity': 'major',
                'description': f"Operating non-type approved exciter: {exciter_model}",
                'equipment_type': 'exciter',
                'equipment_model': exciter_model
            })
        
        # Check amplifier
        amplifier_model = f"{self.inspection.amplifier_manufacturer or ''} {self.inspection.amplifier_model_number or ''}".strip()
        if amplifier_model in non_approved_equipment:
            violations.append({
                'type': 'TYPE_APPROVAL_VIOLATION',
                'severity': 'major',
                'description': f"Operating non-type approved amplifier: {amplifier_model}",
                'equipment_type': 'amplifier',
                'equipment_model': amplifier_model
            })
        
        return violations
    
    def _check_safety_violations(self) -> List[Dict[str, Any]]:
        """Check for safety-related violations"""
        violations = []
        
        # Check lightning protection
        if not self.inspection.has_lightning_protection:
            violations.append({
                'type': 'SAFETY_VIOLATION',
                'severity': 'minor',
                'description': "Lightning protection not provided",
                'safety_category': 'lightning_protection'
            })
        
        # Check grounding
        if not self.inspection.is_electrically_grounded:
            violations.append({
                'type': 'SAFETY_VIOLATION',
                'severity': 'minor',
                'description': "Tower not electrically grounded",
                'safety_category': 'grounding'
            })
        
        # Check aviation warning lights
        tower_height = float(self.inspection.height_above_ground or 0)
        if tower_height > 60 and not self.inspection.has_aviation_warning_light:
            violations.append({
                'type': 'SAFETY_VIOLATION',
                'severity': 'major',
                'description': f"Aviation warning light required for tower height {tower_height}m",
                'safety_category': 'aviation_warning'
            })
        
        return violations